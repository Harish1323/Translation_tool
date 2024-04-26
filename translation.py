# stream lit
# working model

import streamlit as st
from transformers import MarianMTModel, MarianTokenizer
from docx import Document 
import os
import tempfile
import math

os.environ["PROTOCOL_BUFFERS_PYTHON_IMPLEMENTATION"] = "python"

# Mapping of full language names to language codes
language_mapping = {
    "Japanese": "jap",
    "Spanish": "es",
    "German": "de",
    "French": "fr",
    "Italian": "it"
}

# Function to translate text
def translate_text(text, target_language):
    # Load pre-trained translation model and tokenizer
    model_name = f'Helsinki-NLP/opus-mt-en-{target_language}'
    model = MarianMTModel.from_pretrained(model_name)
    tokenizer = MarianTokenizer.from_pretrained(model_name)

    # Tokenize input text
    inputs = tokenizer(text, return_tensors="pt", padding=True, truncation=True)

    # Translate text
    translated_ids = model.generate(inputs['input_ids'])
    translated_text = tokenizer.batch_decode(translated_ids, skip_special_tokens=True)[0]

    return translated_text

# Function to handle translation from Word file
def handle_translation_from_word(word_file, selected_languages, download_location):
    # Load the input document
    input_doc = Document(word_file)
    
    # Create a dictionary to store translated documents for each language
    translated_docs = {language: Document() for language in selected_languages}
    
    # Iterate through each paragraph in the input document
    for paragraph in input_doc.paragraphs:
        text = paragraph.text
        
        # Translate the text for each selected language
        for language in selected_languages:
            # Translate the paragraph text
            translated_text = translate_text(text, language_mapping[language])
            
            # Add the translated text as a new paragraph in the translated document for the language
            translated_docs[language].add_paragraph(translated_text)
    
    # Save each translated document to the specified download location
    output_files = {}
    for language, translated_doc in translated_docs.items():
        # Specify the output file path for the translated document
        output_file_path = os.path.join(download_location, f"translated_{language}.docx")
        # Save the translated document
        translated_doc.save(output_file_path)
        # Add the file path to the dictionary for tracking
        output_files[language] = output_file_path
    
    # Return the paths to the translated files
    return output_files



# Streamlit app
def main():
    st.title("Language Translator")

    # Language selection
    st.sidebar.title("Select Languages")
    selected_languages = st.sidebar.multiselect("Select languages:", list(language_mapping.keys()))

    # File uploader for Word file
    word_file = st.file_uploader("Upload Word file for translation:", type=["docx"])

    # Download location selection
    download_location = st.text_input("Enter download location:", value=os.getcwd())

    # Translate button
    if st.button("Translate"):
        if not word_file:
            st.warning("Please upload a Word file for translation.")
        elif not selected_languages:
            st.warning("Please select at least one language.")
        else:
            with st.spinner("Translating..."):
                # Call handle_translation_from_word with all required arguments
                translated_files = handle_translation_from_word(word_file, selected_languages, download_location)

            # Notify the user of successful translation and output file locations
            st.success("Translation completed. Output files saved successfully.")
            for language, file_path in translated_files.items():
                st.write(f"Translated {language}: {file_path}")

if __name__ == "__main__":
    main()
